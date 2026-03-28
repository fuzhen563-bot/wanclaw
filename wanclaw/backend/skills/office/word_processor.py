"""
Word文档处理技能 V2.0
基于 python-docx 实现真实文件操作
"""

import os
import logging
from typing import Dict, List, Any
from datetime import datetime

from wanclaw.backend.skills import BaseSkill, SkillResult, SkillCategory, SkillLevel

logger = logging.getLogger(__name__)


class WordProcessorSkill(BaseSkill):
    
    def __init__(self):
        super().__init__()
        self.name = "WordProcessor"
        self.description = "Word处理：创建文档、读取内容、批量替换、提取文本、添加水印"
        self.category = SkillCategory.OFFICE
        self.level = SkillLevel.INTERMEDIATE
        self.required_params = ["action"]
        self.optional_params = {
            "file_path": str, "output_path": str, "content": str,
            "find_text": str, "replace_text": str, "table_data": list,
            "headings": list, "style": str
        }

    async def execute(self, params: Dict[str, Any]) -> SkillResult:
        action = params.get("action", "").lower()
        try:
            method = getattr(self, f"_action_{action}", None)
            if not method:
                return SkillResult(success=False, message=f"不支持的操作: {action}")
            return await method(params)
        except ImportError as e:
            return SkillResult(success=False, message=str(e), error=str(e))
        except Exception as e:
            logger.error(f"Word处理失败: {action} - {e}")
            return SkillResult(success=False, message=f"Word处理失败: {str(e)}", error=str(e))

    async def _load_docx(self):
        try:
            import docx
            return docx
        except ImportError:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    async def _action_create(self, params: Dict) -> SkillResult:
        content = params.get("content", "")
        output_path = params.get("output_path", "document.docx")
        headings = params.get("headings", [])
        docx = await self._load_docx()

        doc = docx.Document()
        doc.core_properties.title = "WanClaw 生成的文档"

        for heading in headings:
            doc.add_heading(heading, level=1)

        if content:
            for para in content.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        doc.save(output_path)
        return SkillResult(
            success=True, message=f"文档创建完成: {output_path}",
            data={"output_file": output_path, "paragraphs": len(doc.paragraphs)}
        )

    async def _action_read(self, params: Dict) -> SkillResult:
        file_path = params.get("file_path", "")

        if not file_path or not os.path.exists(file_path):
            return SkillResult(success=False, message="文件不存在", error="file not found")

        docx = await self._load_docx()
        doc = docx.Document(file_path)

        full_text = []
        tables_text = []
        paragraphs_count = len(doc.paragraphs)
        tables_count = len(doc.tables)

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(row_text)

        return SkillResult(
            success=True, message=f"读取完成，共 {paragraphs_count} 段，{tables_count} 个表格",
            data={
                "source_file": file_path,
                "paragraphs": paragraphs_count,
                "tables": tables_count,
                "text": "\n\n".join(full_text),
                "table_preview": tables_text[:5],
                "word_count": sum(len(t) for t in full_text)
            }
        )

    async def _action_replace(self, params: Dict) -> SkillResult:
        file_path = params.get("file_path", "")
        find_text = params.get("find_text", "")
        replace_text = params.get("replace_text", "")
        output_path = params.get("output_path", file_path.replace(".docx", "_replaced.docx"))

        if not file_path or not os.path.exists(file_path):
            return SkillResult(success=False, message="文件不存在", error="file not found")

        docx = await self._load_docx()
        doc = docx.Document(file_path)
        replacements = 0

        for para in doc.paragraphs:
            if find_text in para.text:
                inline = para.runs
                for run in inline:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        replacements += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find_text in para.text:
                            for run in para.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
                                    replacements += 1

        doc.save(output_path)
        return SkillResult(
            success=True, message=f"替换完成，共替换 {replacements} 处",
            data={"find_text": find_text, "replace_text": replace_text,
                  "replacements_made": replacements, "output_file": output_path}
        )

    async def _action_table(self, params: Dict) -> SkillResult:
        table_data = params.get("table_data", [])
        headers = params.get("headers", [])
        output_path = params.get("output_path", "table.docx")

        docx = await self._load_docx()
        doc = docx.Document()

        table = doc.add_table(rows=0, cols=len(headers) if headers else (len(table_data[0]) if table_data else 0))
        table.style = "Light Grid Accent 1"

        if headers:
            header_row = table.add_row()
            for i, h in enumerate(headers):
                header_row.cells[i].text = str(h)

        for row_data in table_data:
            row = table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val)

        doc.save(output_path)
        return SkillResult(
            success=True, message=f"表格创建完成: {output_path}",
            data={"output_file": output_path, "rows": len(table_data), "cols": len(headers or [])}
        )

    async def _action_extract_text(self, params: Dict) -> SkillResult:
        file_path = params.get("file_path", "")

        if not file_path or not os.path.exists(file_path):
            return SkillResult(success=False, message="文件不存在", error="file not found")

        docx = await self._load_docx()
        doc = docx.Document(file_path)

        all_text = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n".join(all_text)

        return SkillResult(
            success=True, message="文本提取完成",
            data={"source_file": file_path, "extracted_text": full_text,
                  "text_length": len(full_text), "paragraph_count": len(all_text)}
        )
